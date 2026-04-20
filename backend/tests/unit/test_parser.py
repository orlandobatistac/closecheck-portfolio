"""Unit tests for text extraction (Day 2)."""
import io
from pathlib import Path

import pytest


# ── helpers ────────────────────────────────────────────────────────────────────

def _make_docx(text: str, tmp_path: Path) -> Path:
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    dest = tmp_path / "sample.docx"
    doc.save(str(dest))
    return dest


def _make_pdf(text: str, tmp_path: Path) -> Path:
    """Create a minimal single-page PDF with reportlab or fpdf2 if available,
    otherwise skip the test."""
    dest = tmp_path / "sample.pdf"
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, text)
        pdf.output(str(dest))
        return dest
    except ImportError:
        pass
    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(dest))
        c.drawString(50, 750, text)
        c.save()
        return dest
    except ImportError:
        pytest.skip("No PDF-generation library available (fpdf2 or reportlab)")


# ── tests ──────────────────────────────────────────────────────────────────────

def test_extract_docx(tmp_path):
    from app.services.parser import extract_text
    expected = "Buyer: John Smith"
    path = _make_docx(expected, tmp_path)
    result = extract_text(path)
    assert expected in result


def test_extract_docx_multiline(tmp_path):
    from app.services.parser import extract_text
    from docx import Document
    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("Line two")
    dest = tmp_path / "multi.docx"
    doc.save(str(dest))
    result = extract_text(dest)
    assert "Line one" in result
    assert "Line two" in result


def test_extract_pdf(tmp_path):
    from app.services.parser import extract_text
    expected = "Purchase Price 500000"
    path = _make_pdf(expected, tmp_path)
    result = extract_text(path)
    assert expected in result


def test_unsupported_extension(tmp_path):
    from app.services.parser import extract_text
    # Use a file whose type cannot be determined (magic bytes + extension both unknown)
    bad = tmp_path / "file.xyz123unknown"
    bad.write_bytes(b"\x00\x01\x02\x03 random binary content with no known signature")
    with pytest.raises(ValueError, match="Unsupported file type|Cannot determine format"):
        extract_text(bad)


def test_empty_docx(tmp_path):
    from app.services.parser import extract_text
    from docx import Document
    doc = Document()
    dest = tmp_path / "empty.docx"
    doc.save(str(dest))
    result = extract_text(dest)
    assert result == ""
