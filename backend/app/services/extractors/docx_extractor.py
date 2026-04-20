"""
DOCX extractor — paragraphs + tables via python-docx.
"""
import logging
from pathlib import Path

from app.services.extractors import BaseExtractor, ParsedDocument, compute_sha256

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):

    def extract(self, path: Path) -> ParsedDocument:
        warnings: list[str] = []
        chunks: list[str] = []
        table_count = 0
        paragraph_count = 0

        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            return ParsedDocument(
                filename=path.name,
                file_type="docx",
                text="",
                extraction_method="failed",
                warnings=["python-docx is not installed."],
                sha256=compute_sha256(path),
            )

        try:
            doc = Document(str(path))
        except PackageNotFoundError:
            return ParsedDocument(
                filename=path.name,
                file_type="docx",
                text="",
                extraction_method="failed",
                warnings=["Corrupt or invalid DOCX file."],
                sha256=compute_sha256(path),
            )
        except Exception as exc:  # noqa: BLE001
            return ParsedDocument(
                filename=path.name,
                file_type="docx",
                text="",
                extraction_method="failed",
                warnings=[f"DOCX open error: {exc}"],
                sha256=compute_sha256(path),
            )

        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                chunks.append(stripped)
                paragraph_count += 1

        for table in doc.tables:
            table_count += 1
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cells))
            chunks.append("\n".join(rows))

        return ParsedDocument(
            filename=path.name,
            file_type="docx",
            text="\n".join(chunks).strip(),
            metadata={"paragraph_count": paragraph_count, "table_count": table_count},
            extraction_method="native",
            warnings=warnings,
            sha256=compute_sha256(path),
        )
