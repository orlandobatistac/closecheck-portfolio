"""
Generate a comprehensive test dataset for CloseCheck pipeline validation.

Transaction scenario:
  Property:       742 Evergreen Terrace, Charlotte, NC 28201  (Parcel 217-541-09)
  Buyer (correct):  James R. Mitchell and Sarah L. Mitchell
  Buyer (wrong):    Jim Mitchell  → triggers PA-001 (name mismatch)
  Seller:           Robert D. Greenwood and Linda K. Greenwood
  Lender:           Piedmont Mortgage Partners, LLC
  Title Co:         Carolina Closing Services, Inc.
  Purchase Price (correct): $342,500.00   (wrong: $345,000.00  → PA-003)
  Seller Credits (correct): $3,500.00     (wrong: $5,000.00    → CD-003)
  Closing Date (correct):   June 15, 2026 (wrong: July 1, 2026 → PA-005)
  Loan Amount:    $314,050.00
  Cash to Close:  $34,892.50
  Interest Rate:  6.875% Fixed 30-Year
  Parcel:         217-541-09

Output structure:
  sample-docs/test_dataset/
  ├── generate_dataset.py
  ├── manifest.json
  └── documents/
      ├── normal/        11 files
      ├── problematic/   13 files
      └── archives/       3 ZIP files

Run:
  py sample-docs/test_dataset/generate_dataset.py
"""

from __future__ import annotations

import csv
import io
import json
import shutil
import zipfile
from pathlib import Path

# ── reportlab ────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt

# ── openpyxl ──────────────────────────────────────────────────────────────────
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ── Pillow ────────────────────────────────────────────────────────────────────
from PIL import Image, ImageDraw, ImageFilter, ImageFont

# =============================================================================
# CONSTANTS — transaction data
# =============================================================================

PROPERTY_ADDRESS  = "742 Evergreen Terrace, Charlotte, NC 28201"
PARCEL            = "217-541-09"
LEGAL_DESCRIPTION = "Lot 7, Block 12, Pinewood Estates Subdivision, Mecklenburg County, NC"

BUYER_CORRECT     = "James R. Mitchell and Sarah L. Mitchell"
BUYER_WRONG       = "Jim Mitchell"            # triggers PA-001
SELLER            = "Robert D. Greenwood and Linda K. Greenwood"
LENDER            = "Piedmont Mortgage Partners, LLC"
TITLE_CO          = "Carolina Closing Services, Inc."

PURCHASE_PRICE_CORRECT = "$342,500.00"
PURCHASE_PRICE_WRONG   = "$345,000.00"        # triggers PA-003

SELLER_CREDITS_CORRECT = "$3,500.00"
SELLER_CREDITS_WRONG   = "$5,000.00"          # triggers CD-003

CLOSING_DATE_CORRECT   = "June 15, 2026"
CLOSING_DATE_WRONG     = "July 1, 2026"       # triggers PA-005

LOAN_AMOUNT   = "$314,050.00"
CASH_TO_CLOSE = "$34,892.50"
INTEREST_RATE = "6.875%"
LOAN_TYPE     = "Conventional 30-Year Fixed"

HOA_NAME    = "Pinewood Estates Homeowners Association"
INS_COMPANY = "Carolinas Home Insurance Group"
INS_POLICY  = "HOP-2026-88431"

ISSUE_DATE  = "April 19, 2026"

# =============================================================================
# DIRECTORIES
# =============================================================================

BASE     = Path(__file__).parent
DOCS     = BASE / "documents"
NORMAL   = DOCS / "normal"
PROBLEM  = DOCS / "problematic"
ARCHIVES = DOCS / "archives"


def _mkdir():
    for d in (NORMAL, PROBLEM, ARCHIVES):
        d.mkdir(parents=True, exist_ok=True)


# =============================================================================
# REPORTLAB HELPERS
# =============================================================================

def _styles():
    s = getSampleStyleSheet()
    h1   = ParagraphStyle("H1",   parent=s["Heading1"], fontSize=14, spaceAfter=12)
    h2   = ParagraphStyle("H2",   parent=s["Heading2"], fontSize=11, spaceAfter=8)
    body = ParagraphStyle("Body", parent=s["Normal"],   fontSize=10, spaceAfter=6, leading=14)
    small = ParagraphStyle("Small", parent=s["Normal"], fontSize=8,  spaceAfter=4, leading=11)
    return h1, h2, body, small


def _build_pdf(out_path: Path, story: list):
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch,  bottomMargin=inch,
    )
    doc.build(story)
    print(f"  Created: {out_path.relative_to(BASE)}")


def _p(text: str, style) -> Paragraph:
    return Paragraph(text, style)


def _sp(height=0.1) -> Spacer:
    return Spacer(1, height * inch)


def _tbl(data: list[list], col_widths=None) -> Table:
    tbl = Table(data, colWidths=col_widths)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a4b8c")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4ff")]),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tbl


# =============================================================================
# SECTION C — NORMAL DOCUMENTS
# =============================================================================

def make_purchase_agreement():
    """normal/purchase_agreement.pdf — correct buyer, price, date"""
    h1, h2, body, small = _styles()
    story = [
        _p("RESIDENTIAL PURCHASE AGREEMENT", h1),
        _p(f"Effective Date: {ISSUE_DATE}", body),
        _p("MLS#: CLT-2026-04812", body),
        _sp(0.15),
        _p("PARTIES", h2),
        _p(f"Buyer(s):  {BUYER_CORRECT}", body),
        _p(f"Seller(s): {SELLER}", body),
        _sp(0.1),
        _p("PROPERTY", h2),
        _p(f"Address:           {PROPERTY_ADDRESS}", body),
        _p(f"Parcel Number:     {PARCEL}", body),
        _p(f"Legal Description: {LEGAL_DESCRIPTION}", body),
        _sp(0.1),
        _p("FINANCIAL TERMS", h2),
        _p(f"Purchase Price:          {PURCHASE_PRICE_CORRECT}", body),
        _p("Earnest Money Deposit:   $6,850.00", body),
        _p(f"Closing Date:            {CLOSING_DATE_CORRECT}", body),
        _p(f"Financing:               {LOAN_AMOUNT} conventional loan via {LENDER}", body),
        _sp(0.1),
        _p("CONTINGENCIES", h2),
        _p("1. Inspection Contingency — Buyer has 14 days from Effective Date to complete inspections.", body),
        _p("2. Financing Contingency — Subject to Buyer obtaining written loan commitment by May 30, 2026.", body),
        _p("3. Appraisal Contingency — Property must appraise at or above the Purchase Price.", body),
        _p("4. Title Contingency — Title must be delivered free and clear of all encumbrances.", body),
        _sp(0.1),
        _p("INCLUSIONS / EXCLUSIONS", h2),
        _p("Included: All built-in appliances, ceiling fans, window treatments, garage door openers.", body),
        _p("Excluded: Dining room chandelier, backyard playset.", body),
        _sp(0.1),
        _p("CLOSING COSTS", h2),
        _p(f"Seller agrees to contribute {SELLER_CREDITS_CORRECT} toward Buyer's closing costs.", body),
        _p("Buyer and Seller to pay their respective costs as customary in Mecklenburg County.", body),
        _sp(0.1),
        _p("SIGNATURES", h2),
        _p(f"Buyer Signature:  ________________________  Date: {ISSUE_DATE}", body),
        _p(f"Buyer Signature:  ________________________  Date: {ISSUE_DATE}", body),
        _p(f"Seller Signature: ________________________  Date: {ISSUE_DATE}", body),
        _p(f"Seller Signature: ________________________  Date: {ISSUE_DATE}", body),
    ]
    _build_pdf(NORMAL / "purchase_agreement.pdf", story)


def make_closing_disclosure():
    """normal/closing_disclosure.pdf — correct price, credits, date"""
    h1, h2, body, small = _styles()
    story = [
        _p("CLOSING DISCLOSURE", h1),
        _p("This form is a statement of final loan terms and closing costs.", body),
        _p(f"Closing Date: {CLOSING_DATE_CORRECT}    Disbursement Date: {CLOSING_DATE_CORRECT}", body),
        _sp(0.15),
        _p("TRANSACTION INFORMATION", h2),
        _p(f"Borrower:        {BUYER_CORRECT}", body),
        _p(f"Seller:          {SELLER}", body),
        _p(f"Lender:          {LENDER}", body),
        _p(f"Property:        {PROPERTY_ADDRESS}", body),
        _sp(0.1),
        _p("LOAN INFORMATION", h2),
        _p(f"Purchase Price:  {PURCHASE_PRICE_CORRECT}", body),
        _p(f"Loan Amount:     {LOAN_AMOUNT}", body),
        _p(f"Loan Type:       {LOAN_TYPE}", body),
        _p(f"Interest Rate:   {INTEREST_RATE} Fixed", body),
        _p("Loan Term:       30 years", body),
        _sp(0.1),
        _p("CLOSING COST DETAILS", h2),
        _tbl([
            ["Item", "Borrower Pays", "Seller Pays"],
            ["Loan Origination Fee",      "$3,140.50", "—"],
            ["Appraisal Fee",             "$550.00",   "—"],
            ["Title Search",              "$250.00",   "—"],
            ["Lender's Title Insurance",  "$875.00",   "—"],
            ["Owner's Title Insurance",   "—",         "$1,200.00"],
            ["Recording Fees",            "$145.00",   "—"],
            ["Property Tax Proration",    "$620.00",   "—"],
            ["Seller Credit",             f"({SELLER_CREDITS_CORRECT})", "—"],
            ["TOTAL",                     "$4,080.50", "$1,200.00"],
        ], col_widths=[3.2*inch, 1.7*inch, 1.7*inch]),
        _sp(0.1),
        _p("CASH TO CLOSE SUMMARY", h2),
        _p(f"Cash to Close from Borrower: {CASH_TO_CLOSE}", body),
        _p(f"Seller Credits Applied:      {SELLER_CREDITS_CORRECT}", body),
        _sp(0.1),
        _p("CONTACT INFORMATION", h2),
        _p(f"Settlement Agent: {TITLE_CO}", body),
        _p("This is an official Closing Disclosure under TRID regulations.", small),
    ]
    _build_pdf(NORMAL / "closing_disclosure.pdf", story)


def make_title_commitment():
    """normal/title_commitment.pdf"""
    h1, h2, body, small = _styles()
    story = [
        _p("TITLE COMMITMENT (ALTA Form 2016)", h1),
        _p(f"Commitment Date: {ISSUE_DATE}", body),
        _sp(0.15),
        _p("SCHEDULE A — TRANSACTION DETAILS", h2),
        _p(f"Proposed Insured (Buyer):  {BUYER_CORRECT}", body),
        _p(f"Proposed Insured (Lender): {LENDER}", body),
        _p(f"Property Address:          {PROPERTY_ADDRESS}", body),
        _p(f"Legal Description:         {LEGAL_DESCRIPTION}", body),
        _p(f"Parcel Number:             {PARCEL}", body),
        _p(f"Title Insurance Amount:    {PURCHASE_PRICE_CORRECT}", body),
        _p(f"Issuing Agent:             {TITLE_CO}", body),
        _sp(0.1),
        _p("SCHEDULE B — REQUIREMENTS", h2),
        _p("The following must be satisfied before title insurance is issued:", body),
        _p("1. Payment in full of the purchase price.", body),
        _p("2. Execution and recording of Warranty Deed from Seller to Buyer.", body),
        _p("3. Execution and recording of Deed of Trust in favor of Lender.", body),
        _p("4. Release of existing mortgage lien (Account #: PMM-2019-4421) held by BB&T.", body),
        _p("5. Payment of all outstanding property taxes through closing date.", body),
        _sp(0.1),
        _p("SCHEDULE B — EXCEPTIONS", h2),
        _p("The policy does NOT insure against loss arising from:", body),
        _p("1. Current year property taxes not yet due and payable.", body),
        _p("2. Utility easement along the northern property boundary (recorded DB 4421, Pg 187).", body),
        _p("3. Subdivision covenants, conditions, and restrictions per Plat Book 22, Page 14.", body),
        _p("4. Rights of parties in possession not shown in the public records.", body),
        _sp(0.1),
        _p("SCHEDULE C — LIENS AND ENCUMBRANCES", h2),
        _p("Open Mortgage Liens:  1 (to be released at closing)", body),
        _p("Judgments:            None", body),
        _p("Tax Liens:            None", body),
        _p("Mechanic's Liens:     None", body),
        _sp(0.1),
        _p(f"Authorized Agent: ________________________  Date: {ISSUE_DATE}", body),
    ]
    _build_pdf(NORMAL / "title_commitment.pdf", story)


def make_loan_note():
    """normal/loan_note.pdf"""
    h1, h2, body, small = _styles()
    story = [
        _p("PROMISSORY NOTE", h1),
        _p(f"Date: {CLOSING_DATE_CORRECT}", body),
        _p(f"City: Charlotte, State: North Carolina", body),
        _sp(0.15),
        _p("BORROWER INFORMATION", h2),
        _p(f"Borrower(s): {BUYER_CORRECT}", body),
        _p(f"Property Address: {PROPERTY_ADDRESS}", body),
        _sp(0.1),
        _p("LOAN TERMS", h2),
        _p(f"Principal Amount:   {LOAN_AMOUNT}", body),
        _p(f"Interest Rate:      {INTEREST_RATE} Fixed", body),
        _p("Loan Term:          360 months (30 years)", body),
        _p("Monthly P&I:        $2,062.18", body),
        _p(f"First Payment Due:  July 1, 2026", body),
        _p(f"Maturity Date:      June 1, 2056", body),
        _sp(0.1),
        _p("LENDER INFORMATION", h2),
        _p(f"Lender: {LENDER}", body),
        _p("NMLS ID: 1887432", body),
        _p("Address: 1200 Tryon Street, Suite 400, Charlotte, NC 28202", body),
        _sp(0.1),
        _p("PROMISE TO PAY", h2),
        _p(
            "In return for a loan I have received, I promise to pay U.S. $314,050.00 (this amount "
            "is called 'Principal'), plus interest, to the order of the Lender. I will make a "
            "payment every month on the 1st day of each month beginning on July 1, 2026.",
            body,
        ),
        _sp(0.1),
        _p("PREPAYMENT", h2),
        _p("I have the right to make payments of Principal at any time before they are due. "
           "A payment of Principal only is known as a 'Prepayment'. There is no prepayment penalty.", body),
        _sp(0.1),
        _p("SIGNATURES", h2),
        _p(f"Borrower: ________________________  Date: {CLOSING_DATE_CORRECT}", body),
        _p(f"Borrower: ________________________  Date: {CLOSING_DATE_CORRECT}", body),
        _p(f"Lender Representative: ____________  Date: {CLOSING_DATE_CORRECT}", body),
    ]
    _build_pdf(NORMAL / "loan_note.pdf", story)


def make_insurance_binder():
    """normal/insurance_binder.pdf"""
    h1, h2, body, small = _styles()
    story = [
        _p("HOMEOWNER'S INSURANCE BINDER", h1),
        _p(f"Issue Date: {ISSUE_DATE}", body),
        _p(f"Policy Number: {INS_POLICY}", body),
        _sp(0.15),
        _p("INSURED INFORMATION", h2),
        _p(f"Named Insured: {BUYER_CORRECT}", body),
        _p(f"Property Address: {PROPERTY_ADDRESS}", body),
        _sp(0.1),
        _p("COVERAGE DETAILS", h2),
        _tbl([
            ["Coverage Type",        "Limit",        "Deductible"],
            ["Dwelling (Coverage A)", "$370,000",     "$1,000"],
            ["Other Structures (B)",  "$37,000",      "$1,000"],
            ["Personal Property (C)", "$185,000",     "$1,000"],
            ["Loss of Use (D)",       "$74,000",      "—"],
            ["Liability (E)",         "$300,000",     "—"],
            ["Medical Payments (F)",  "$5,000/person","—"],
        ], col_widths=[2.8*inch, 1.8*inch, 1.8*inch]),
        _sp(0.1),
        _p("MORTGAGE CLAUSE", h2),
        _p(f"Mortgagee / Additional Insured: {LENDER}", body),
        _p("Mortgagee Address: 1200 Tryon Street, Suite 400, Charlotte, NC 28202", body),
        _p("ISAOA / ATIMA (Its Successors and Assigns, as their interests may appear)", body),
        _sp(0.1),
        _p("POLICY PERIOD", h2),
        _p(f"Effective: {CLOSING_DATE_CORRECT}  |  Expires: June 15, 2027", body),
        _p(f"Annual Premium: $1,284.00", body),
        _sp(0.1),
        _p(f"Insurance Company: {INS_COMPANY}", body),
        _p(f"Agent: ________________________  License: NC-INS-882341", body),
    ]
    _build_pdf(NORMAL / "insurance_binder.pdf", story)


def make_hoa_certificate():
    """normal/hoa_certificate.pdf"""
    h1, h2, body, small = _styles()
    story = [
        _p("HOMEOWNER ASSOCIATION CERTIFICATE", h1),
        _p(f"Date Issued: {ISSUE_DATE}", body),
        _sp(0.15),
        _p("HOA INFORMATION", h2),
        _p(f"Association Name: {HOA_NAME}", body),
        _p("Management Company: Carolinas Community Management, LLC", body),
        _p("Address: 4000 Westchase Blvd, Suite 200, Charlotte, NC 28217", body),
        _p("Phone: (704) 555-0182  |  Email: info@pinewoodestates-hoa.com", body),
        _sp(0.1),
        _p("PROPERTY DETAILS", h2),
        _p(f"Property Address: {PROPERTY_ADDRESS}", body),
        _p(f"Parcel Number: {PARCEL}", body),
        _p(f"Current Owner (Seller): {SELLER}", body),
        _p(f"Buyer: {BUYER_CORRECT}", body),
        _sp(0.1),
        _p("ASSESSMENT INFORMATION", h2),
        _tbl([
            ["Assessment Type",         "Amount",    "Frequency"],
            ["Regular Monthly Dues",     "$245.00",   "Monthly"],
            ["Capital Reserve Fund",     "$30.00",    "Monthly"],
            ["Total Monthly Assessment", "$275.00",   "Monthly"],
        ], col_widths=[3.0*inch, 1.5*inch, 1.5*inch]),
        _sp(0.1),
        _p("ACCOUNT STATUS", h2),
        _p("Current Balance Due from Seller:  $0.00 (All dues current as of this date)", body),
        _p("No special assessments pending.", body),
        _p("No pending or active litigation against the HOA.", body),
        _sp(0.1),
        _p("TRANSFER INFORMATION", h2),
        _p("Transfer Fee: $350.00 (due at closing, paid by Buyer)", body),
        _p("HOA Documents (CC&Rs, Bylaws, Budget): Available at HOA management office.", body),
        _sp(0.1),
        _p(f"Authorized Representative: ________________________  Date: {ISSUE_DATE}", body),
    ]
    _build_pdf(NORMAL / "hoa_certificate.pdf", story)


def make_wire_instructions():
    """normal/wire_instructions.txt"""
    content = f"""WIRE TRANSFER INSTRUCTIONS
Issued By: {TITLE_CO}
Date: {ISSUE_DATE}
Transaction: {PROPERTY_ADDRESS}

SETTLEMENT AGENT WIRE DETAILS
==============================
Bank Name:          First Citizens Bank & Trust Co.
ABA Routing Number: 053100300
Account Number:     7841209654
Account Name:       {TITLE_CO} — Escrow Trust Account
Reference:          Mitchell/Greenwood — {PARCEL}

FUNDS DUE AT CLOSING
====================
Cash to Close (Buyer):          {CASH_TO_CLOSE}
Buyer Paid-Outside-Closing:     $550.00  (appraisal)
Lender Wire (Loan Proceeds):    {LOAN_AMOUNT}

IMPORTANT WIRE INSTRUCTIONS
============================
1. Wire funds must be received ONE BUSINESS DAY before closing.
2. Wires received after 3:00 PM ET will NOT be posted until the following business day.
3. ALWAYS call to verify wire details before sending: (704) 555-0199.
4. Carolinas Closing Services will NEVER change wire instructions via email.
5. Confirm any changes to this document by calling our office directly.
6. If you suspect fraud, contact your bank immediately and do NOT send the wire.

CONTACT
=======
Closing Officer: Margaret J. Hollowell
Direct Phone:    (704) 555-0199
Email:           mhollowell@carolinasclosing.com

This document contains confidential financial information.
"""
    path = NORMAL / "wire_instructions.txt"
    path.write_text(content, encoding="utf-8")
    print(f"  Created: {path.relative_to(BASE)}")


def make_buyer_info_json():
    """normal/buyer_info.json"""
    data = {
        "transaction_id": "CC-2026-04-7421",
        "property": {
            "address": PROPERTY_ADDRESS,
            "parcel": PARCEL,
            "legal_description": LEGAL_DESCRIPTION,
            "county": "Mecklenburg",
            "state": "NC",
            "zip": "28201",
        },
        "buyers": [
            {
                "name": "James R. Mitchell",
                "ssn_last4": "XXXX",
                "dob": "1982-03-14",
                "email": "james.mitchell@email.com",
                "phone": "(704) 555-0241",
                "id_type": "NC Driver License",
                "id_number": "NC-****-8841",
            },
            {
                "name": "Sarah L. Mitchell",
                "ssn_last4": "XXXX",
                "dob": "1984-07-29",
                "email": "sarah.mitchell@email.com",
                "phone": "(704) 555-0242",
                "id_type": "NC Driver License",
                "id_number": "NC-****-2217",
            },
        ],
        "sellers": [
            {"name": "Robert D. Greenwood"},
            {"name": "Linda K. Greenwood"},
        ],
        "financial": {
            "purchase_price": 342500.00,
            "loan_amount": 314050.00,
            "cash_to_close": 34892.50,
            "seller_credits": 3500.00,
            "earnest_money": 6850.00,
            "interest_rate": "6.875%",
            "closing_date": CLOSING_DATE_CORRECT,
        },
        "lender": {
            "name": LENDER,
            "nmls": "1887432",
            "contact": "David Hartmann",
            "phone": "(704) 555-0300",
        },
        "title_company": {
            "name": TITLE_CO,
            "contact": "Margaret J. Hollowell",
            "phone": "(704) 555-0199",
        },
    }
    path = NORMAL / "buyer_info.json"
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    print(f"  Created: {path.relative_to(BASE)}")


def make_closing_costs_xlsx():
    """normal/closing_costs.xlsx — Sheet 1: itemized costs, Sheet 2: amortization (first 6 mo)"""
    wb = openpyxl.Workbook()

    # ── Sheet 1: Closing Costs ─────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Closing Costs"

    header_fill = PatternFill("solid", fgColor="1a4b8c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    total_fill  = PatternFill("solid", fgColor="dce8f5")
    total_font  = Font(bold=True, size=10)
    thin_border = Border(
        bottom=Side(style="thin"),
        top=Side(style="thin"),
        left=Side(style="thin"),
        right=Side(style="thin"),
    )

    headers = ["Description", "Paid By", "Amount ($)"]
    ws1.append(headers)
    for col, h in enumerate(headers, start=1):
        cell = ws1.cell(row=1, column=col)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    buyer_items = [
        ("Loan Origination Fee (1%)",      "Buyer",  3140.50),
        ("Appraisal Fee",                  "Buyer",   550.00),
        ("Credit Report",                  "Buyer",    72.00),
        ("Flood Determination",            "Buyer",    12.00),
        ("Title Search",                   "Buyer",   250.00),
        ("Lender's Title Insurance",       "Buyer",   875.00),
        ("Survey Fee",                     "Buyer",   400.00),
        ("Recording Fee — Deed",           "Buyer",    75.00),
        ("Recording Fee — Deed of Trust",  "Buyer",    70.00),
        ("Property Tax Proration",         "Buyer",   620.00),
        ("Homeowner's Insurance (1st yr)", "Buyer",  1284.00),
        ("HOA Transfer Fee",               "Buyer",   350.00),
        ("Seller Credit",                  "Buyer", -3500.00),
    ]
    seller_items = [
        ("Owner's Title Insurance",        "Seller", 1200.00),
        ("Real Estate Commission (5.5%)",  "Seller", 18837.50),
        ("Attorney Fee",                   "Seller",  650.00),
        ("Recording Fee — Mortgage Release","Seller",  35.00),
    ]

    buyer_total  = sum(a for _, _, a in buyer_items)
    seller_total = sum(a for _, _, a in seller_items)

    for row in buyer_items + seller_items:
        ws1.append(list(row))

    ws1.append(["TOTAL", "", ""])
    ws1.append(["  Buyer Cash to Close", "Buyer",  buyer_total])
    ws1.append(["  Seller Net Proceeds", "Seller", seller_total])

    # Style totals
    for row_num in [len(buyer_items) + len(seller_items) + 2,
                    len(buyer_items) + len(seller_items) + 3,
                    len(buyer_items) + len(seller_items) + 4]:
        for col in range(1, 4):
            c = ws1.cell(row=row_num, column=col)
            c.fill = total_fill
            c.font = total_font

    ws1.column_dimensions["A"].width = 38
    ws1.column_dimensions["B"].width = 10
    ws1.column_dimensions["C"].width = 14

    # ── Sheet 2: Amortization Schedule (first 12 months) ──────────
    ws2 = wb.create_sheet("Amortization")
    ws2.append(["Month", "Payment", "Principal", "Interest", "Balance"])
    hdr_row = ws2[1]
    for cell in hdr_row:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    principal = 314050.00
    rate_monthly = 0.06875 / 12
    payment = 2062.18
    for month in range(1, 13):
        interest = round(principal * rate_monthly, 2)
        principal_paid = round(payment - interest, 2)
        principal = round(principal - principal_paid, 2)
        ws2.append([month, payment, principal_paid, interest, principal])

    for col in ["A", "B", "C", "D", "E"]:
        ws2.column_dimensions[col].width = 14

    path = NORMAL / "closing_costs.xlsx"
    wb.save(str(path))
    print(f"  Created: {path.relative_to(BASE)}")


def make_settlement_summary_csv():
    """normal/settlement_summary.csv — cross-document comparison table"""
    rows = [
        ["Field", "Purchase Agreement", "Closing Disclosure", "Loan Note", "Status"],
        ["Buyer Name",       BUYER_CORRECT,           BUYER_CORRECT,           BUYER_CORRECT,     "MATCH"],
        ["Seller Name",      SELLER,                  SELLER,                  "N/A",             "MATCH"],
        ["Property Address", PROPERTY_ADDRESS,         PROPERTY_ADDRESS,        PROPERTY_ADDRESS,  "MATCH"],
        ["Purchase Price",   PURCHASE_PRICE_CORRECT,  PURCHASE_PRICE_CORRECT,  "N/A",             "MATCH"],
        ["Loan Amount",      LOAN_AMOUNT,             LOAN_AMOUNT,             LOAN_AMOUNT,       "MATCH"],
        ["Closing Date",     CLOSING_DATE_CORRECT,    CLOSING_DATE_CORRECT,    "N/A",             "MATCH"],
        ["Seller Credits",   SELLER_CREDITS_CORRECT,  SELLER_CREDITS_CORRECT,  "N/A",             "MATCH"],
        ["Cash to Close",    CASH_TO_CLOSE,           CASH_TO_CLOSE,           "N/A",             "MATCH"],
        ["Interest Rate",    INTEREST_RATE,           INTEREST_RATE,           INTEREST_RATE,     "MATCH"],
        ["Lender",           LENDER,                  LENDER,                  LENDER,            "MATCH"],
    ]
    path = NORMAL / "settlement_summary.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"  Created: {path.relative_to(BASE)}")


def make_property_photo():
    """normal/property_photo.jpg — synthetic house silhouette via Pillow"""
    W, H = 800, 600
    img = Image.new("RGB", (W, H), color=(135, 180, 100))  # green lawn
    draw = ImageDraw.Draw(img)

    # Sky
    draw.rectangle([(0, 0), (W, H // 2)], fill=(100, 160, 230))

    # House body
    draw.rectangle([(150, 280), (650, 500)], fill=(220, 200, 170), outline=(80, 60, 40), width=3)

    # Roof (triangle)
    draw.polygon([(120, 280), (400, 100), (680, 280)], fill=(140, 80, 50), outline=(80, 40, 20), width=2)

    # Door
    draw.rectangle([(350, 400), (450, 500)], fill=(100, 60, 30), outline=(50, 30, 10), width=2)
    draw.ellipse([(430, 445), (445, 460)], fill=(220, 180, 0))  # doorknob

    # Windows (left)
    draw.rectangle([(190, 320), (280, 390)], fill=(180, 220, 255), outline=(80, 60, 40), width=2)
    draw.line([(235, 320), (235, 390)], fill=(80, 60, 40), width=1)
    draw.line([(190, 355), (280, 355)], fill=(80, 60, 40), width=1)

    # Windows (right)
    draw.rectangle([(520, 320), (610, 390)], fill=(180, 220, 255), outline=(80, 60, 40), width=2)
    draw.line([(565, 320), (565, 390)], fill=(80, 60, 40), width=1)
    draw.line([(520, 355), (610, 355)], fill=(80, 60, 40), width=1)

    # Chimney
    draw.rectangle([(540, 80), (590, 200)], fill=(160, 100, 70), outline=(80, 40, 20), width=2)

    # Walkway
    draw.polygon([(370, 500), (430, 500), (460, 580), (340, 580)], fill=(180, 170, 160))

    # Lawn
    draw.rectangle([(0, 500), (W, H)], fill=(100, 155, 80))

    # Address sign
    draw.rectangle([(310, 530), (490, 560)], fill=(240, 235, 220), outline=(80, 60, 40), width=1)
    try:
        font = ImageFont.truetype("arial.ttf", 16)
    except Exception:
        font = ImageFont.load_default()
    draw.text((320, 534), "742 Evergreen Terrace", fill=(40, 40, 40), font=font)

    path = NORMAL / "property_photo.jpg"
    img.save(str(path), "JPEG", quality=92)
    print(f"  Created: {path.relative_to(BASE)}")
    return path  # returned so duplicate can reuse bytes


# =============================================================================
# SECTION D — PROBLEMATIC DOCUMENTS
# =============================================================================

def make_cd_price_conflict():
    """problematic/cd_amended_price_conflict.pdf — price=$345k (PA-003) + credits=$5k (CD-003)"""
    h1, h2, body, small = _styles()
    story = [
        _p("CLOSING DISCLOSURE — AMENDED VERSION", h1),
        _p("NOTICE: This is a revised CD reflecting amended contract terms.", body),
        _p(f"Closing Date: {CLOSING_DATE_CORRECT}", body),
        _sp(0.15),
        _p("TRANSACTION INFORMATION", h2),
        _p(f"Borrower: {BUYER_CORRECT}", body),
        _p(f"Seller:   {SELLER}", body),
        _p(f"Lender:   {LENDER}", body),
        _p(f"Property: {PROPERTY_ADDRESS}", body),
        _sp(0.1),
        _p("LOAN INFORMATION", h2),
        _p(f"Purchase Price:  {PURCHASE_PRICE_WRONG}  ← AMENDED (orig. {PURCHASE_PRICE_CORRECT})", body),
        _p(f"Loan Amount:     {LOAN_AMOUNT}", body),
        _p(f"Interest Rate:   {INTEREST_RATE} Fixed", body),
        _sp(0.1),
        _p("CLOSING COST DETAILS", h2),
        _tbl([
            ["Item",                      "Borrower Pays",              "Seller Pays"],
            ["Loan Origination Fee",      "$3,140.50",                  "—"],
            ["Appraisal Fee",             "$550.00",                    "—"],
            ["Seller Credit",             f"({SELLER_CREDITS_WRONG})",  "—"],
            ["TOTAL",                     "$2,330.50",                  "$1,200.00"],
        ], col_widths=[3.2*inch, 1.7*inch, 1.7*inch]),
        _sp(0.1),
        _p(f"Cash to Close from Borrower: {CASH_TO_CLOSE}", body),
        _p(f"Seller Credits Applied:      {SELLER_CREDITS_WRONG}", body),
        _sp(0.1),
        _p("⚠  This document contains values that differ from the original Purchase Agreement.", small),
    ]
    _build_pdf(PROBLEM / "cd_amended_price_conflict.pdf", story)


def make_pa_signed_name_mismatch():
    """problematic/pa_signed_name_mismatch.docx — buyer signed as 'Jim Mitchell' (PA-001)"""
    doc = DocxDocument()
    doc.add_heading("RESIDENTIAL PURCHASE AGREEMENT — SIGNED COPY", 0)
    doc.add_paragraph(f"Effective Date: {ISSUE_DATE}")

    doc.add_heading("PARTIES", 1)
    doc.add_paragraph(f"Buyer(s):  {BUYER_CORRECT}")
    doc.add_paragraph(f"Seller(s): {SELLER}")

    doc.add_heading("PROPERTY", 1)
    doc.add_paragraph(f"Address: {PROPERTY_ADDRESS}")
    doc.add_paragraph(f"Parcel: {PARCEL}")

    doc.add_heading("FINANCIAL TERMS", 1)
    doc.add_paragraph(f"Purchase Price: {PURCHASE_PRICE_CORRECT}")
    doc.add_paragraph(f"Closing Date:   {CLOSING_DATE_CORRECT}")
    doc.add_paragraph(f"Seller Credits: {SELLER_CREDITS_CORRECT}")

    doc.add_heading("CONTINGENCIES", 1)
    doc.add_paragraph("1. Inspection — 14 days from effective date.")
    doc.add_paragraph("2. Financing — Written commitment by May 30, 2026.")
    doc.add_paragraph("3. Appraisal — At or above purchase price.")

    doc.add_heading("SIGNATURES", 1)
    doc.add_paragraph("Typed/Printed Buyer Name: James R. Mitchell and Sarah L. Mitchell")
    doc.add_paragraph("")

    # The mismatch: signed as "Jim Mitchell" instead of "James R. Mitchell"
    p = doc.add_paragraph()
    run = p.add_run("Buyer Signature (as signed): Jim Mitchell")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(f"Buyer Signature: Sarah L. Mitchell")
    doc.add_paragraph(f"Seller Signature: Robert D. Greenwood")
    doc.add_paragraph(f"Seller Signature: Linda K. Greenwood")
    doc.add_paragraph(f"Date: {ISSUE_DATE}")

    path = PROBLEM / "pa_signed_name_mismatch.docx"
    doc.save(str(path))
    print(f"  Created: {path.relative_to(BASE)}")


def make_loan_estimate_date_conflict():
    """problematic/loan_estimate_date_conflict.pdf — closing date = July 1, 2026 (PA-005)"""
    h1, h2, body, small = _styles()
    story = [
        _p("LOAN ESTIMATE", h1),
        _p("This form is an estimate of the loan terms and closing costs.", body),
        _sp(0.15),
        _p("LOAN TERMS", h2),
        _p(f"Loan Amount:     {LOAN_AMOUNT}", body),
        _p(f"Interest Rate:   {INTEREST_RATE} Fixed", body),
        _p(f"Loan Type:       {LOAN_TYPE}", body),
        _sp(0.1),
        _p("TRANSACTION INFORMATION", h2),
        _p(f"Borrower:        {BUYER_CORRECT}", body),
        _p(f"Property:        {PROPERTY_ADDRESS}", body),
        _p(f"Purchase Price:  {PURCHASE_PRICE_CORRECT}", body),
        # Intentional date conflict
        _p(f"Estimated Closing Date:  {CLOSING_DATE_WRONG}  ← differs from Purchase Agreement", body),
        _p(f"Rate Lock Expires:       July 15, 2026", body),
        _sp(0.1),
        _p("PROJECTED PAYMENTS", h2),
        _tbl([
            ["Payment Component",  "Monthly Amount"],
            ["Principal & Interest", "$2,062.18"],
            ["Mortgage Insurance",   "$0.00"],
            ["Estimated Escrow",     "$520.00"],
            ["Total Monthly Payment","$2,582.18"],
        ], col_widths=[3.5*inch, 3.1*inch]),
        _sp(0.1),
        _p("CLOSING COST ESTIMATE", h2),
        _p("Estimated Closing Costs: $5,330.50", body),
        _p(f"Cash to Close (estimated): {CASH_TO_CLOSE}", body),
        _sp(0.1),
        _p(f"Prepared by: {LENDER}  |  Date: {ISSUE_DATE}", small),
        _p("⚠  Closing date on this estimate does not match the Purchase Agreement.", small),
    ]
    _build_pdf(PROBLEM / "loan_estimate_date_conflict.pdf", story)


def make_corrupt_inspection_report():
    """problematic/corrupt_inspection_report.pdf — valid PDF header + garbage body"""
    # A real PDF header that fitz will open, then garbage bytes that break content streams
    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    garbage = b"\x00\xff\xfe" * 512 + b"1 0 obj\n<< /Type /Catalog >>\nendobj\n"
    # No xref / startxref → PyMuPDF opens it but cannot extract text → triggers OCR fallback
    content = header + garbage + b"%%EOF\n"
    path = PROBLEM / "corrupt_inspection_report.pdf"
    path.write_bytes(content)
    print(f"  Created: {path.relative_to(BASE)}")


def make_survey_scan_lowres():
    """problematic/survey_scan_lowres.jpg — heavily degraded scan (JPEG quality=15, blur)"""
    W, H = 400, 300
    img = Image.new("RGB", (W, H), color=(240, 235, 220))
    draw = ImageDraw.Draw(img)

    # Simulated survey plat
    draw.rectangle([(40, 40), (360, 260)], outline=(50, 50, 50), width=2)
    draw.line([(40, 150), (360, 150)], fill=(80, 80, 80), width=1)
    draw.line([(200, 40), (200, 260)], fill=(80, 80, 80), width=1)

    try:
        font_sm = ImageFont.truetype("arial.ttf", 12)
    except Exception:
        font_sm = ImageFont.load_default()

    draw.text((60, 55),  "SURVEY PLAT",              fill=(30, 30, 30), font=font_sm)
    draw.text((60, 75),  f"742 Evergreen Terr.",      fill=(30, 30, 30), font=font_sm)
    draw.text((60, 95),  "Mecklenburg Co., NC",       fill=(30, 30, 30), font=font_sm)
    draw.text((60, 115), f"Parcel: {PARCEL}",         fill=(30, 30, 30), font=font_sm)
    draw.text((60, 155), "Lot area: 0.312 ac",        fill=(30, 30, 30), font=font_sm)
    draw.text((60, 175), "N: 35.221042",              fill=(30, 30, 30), font=font_sm)
    draw.text((60, 195), "W: -80.843211",             fill=(30, 30, 30), font=font_sm)
    draw.text((220, 55), "Scale: 1\" = 30'",          fill=(30, 30, 30), font=font_sm)
    draw.text((220, 75), "Surveyor: J. Perkins PLS",  fill=(30, 30, 30), font=font_sm)

    # Shrink to simulate low resolution then upscale
    thumb = img.resize((80, 60), Image.NEAREST).resize((W, H), Image.NEAREST)
    # Add blur and noise
    blurred = thumb.filter(ImageFilter.GaussianBlur(radius=4))

    path = PROBLEM / "survey_scan_lowres.jpg"
    blurred.save(str(path), "JPEG", quality=15)
    print(f"  Created: {path.relative_to(BASE)}")


def make_title_partial():
    """problematic/title_partial_missing_page.pdf — truncated mid-page (simulates missing p.2)"""
    h1, h2, body, small = _styles()
    story = [
        _p("TITLE COMMITMENT (ALTA Form 2016)", h1),
        _p(f"Commitment Date: {ISSUE_DATE}", body),
        _sp(0.15),
        _p("SCHEDULE A — TRANSACTION DETAILS", h2),
        _p(f"Proposed Insured: {BUYER_CORRECT}", body),
        _p(f"Property Address: {PROPERTY_ADDRESS}", body),
        _p(f"Title Insurance Amount: {PURCHASE_PRICE_CORRECT}", body),
        _sp(0.1),
        _p("SCHEDULE B — REQUIREMENTS", h2),
        _p("The following requirements must be met:", body),
        _p("1. Execution and delivery of Warranty Deed from Grantor to Grantee.", body),
        _p("2. Payment in full of the purchase price.", body),
        _p("3. Release of existing lien recorded in Deed Book 8821, Page 447.", body),
        _p("4. Proof of payment of all property taxes through", body),  # truncated sentence
        # Page 2 is intentionally absent — document ends abruptly
        _p("[DOCUMENT TRUNCATED — Page 2 of 4 missing]", small),
    ]
    _build_pdf(PROBLEM / "title_partial_missing_page.pdf", story)


def make_closing_costs_bad_math():
    """problematic/closing_costs_bad_math.xlsx — TOTAL row ≠ SUM of line items"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Closing Costs"

    header_fill = PatternFill("solid", fgColor="1a4b8c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    error_fill  = PatternFill("solid", fgColor="FFD0D0")
    error_font  = Font(bold=True, color="CC0000", size=10)

    ws.append(["Description", "Paid By", "Amount ($)"])
    for col, _ in enumerate(["Description", "Paid By", "Amount ($)"], start=1):
        c = ws.cell(row=1, column=col)
        c.font = header_font
        c.fill = header_fill

    items = [
        ("Loan Origination Fee (1%)",     "Buyer",   3140.50),
        ("Appraisal Fee",                  "Buyer",    550.00),
        ("Credit Report",                  "Buyer",     72.00),
        ("Title Search",                   "Buyer",    250.00),
        ("Lender's Title Insurance",       "Buyer",    875.00),
        ("Survey Fee",                     "Buyer",    400.00),
        ("Recording Fee",                  "Buyer",    145.00),
        ("Property Tax Proration",         "Buyer",    620.00),
        ("Homeowner's Insurance (1st yr)", "Buyer",   1284.00),
        ("HOA Transfer Fee",               "Buyer",    350.00),
        ("Seller Credit",                  "Buyer",  -3500.00),
    ]
    for row in items:
        ws.append(list(row))

    real_total = sum(a for _, _, a in items)  # 4186.50
    wrong_total = real_total + 1250.00  # deliberate $1,250 error

    ws.append(["TOTAL (BUYER CASH TO CLOSE)", "Buyer", wrong_total])

    # Highlight the wrong total in red
    last_row = ws.max_row
    for col in range(1, 4):
        c = ws.cell(row=last_row, column=col)
        c.fill = error_fill
        c.font = error_font

    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 10
    ws.column_dimensions["C"].width = 16

    # Add a note
    ws.cell(row=last_row + 2, column=1).value = (
        f"NOTE: Real sum = ${real_total:,.2f}  |  Total shown = ${wrong_total:,.2f}  "
        f"(difference: $1,250.00 — calculation error)"
    )

    path = PROBLEM / "closing_costs_bad_math.xlsx"
    wb.save(str(path))
    print(f"  Created: {path.relative_to(BASE)}")


def make_addendum_empty_fields():
    """problematic/addendum_empty_fields.docx — unfilled placeholder fields"""
    doc = DocxDocument()
    doc.add_heading("PURCHASE AGREEMENT ADDENDUM", 0)
    doc.add_paragraph(f"This Addendum is attached to and incorporated into the Purchase Agreement")
    doc.add_paragraph(f"dated __________ for the property located at {PROPERTY_ADDRESS}.")
    doc.add_paragraph("")

    doc.add_heading("ADDITIONAL TERMS", 1)

    # Unfilled placeholders
    fields = [
        ("Addendum Issue Date",           "__________"),
        ("Buyer Name(s)",                 "__________"),
        ("Seller Name(s)",                "__________"),
        ("Modified Purchase Price",       "$__________"),
        ("New Closing Date",              "__________"),
        ("Reason for Amendment",          "__________________________________________"),
        ("Special Conditions",            "__________________________________________"),
        ("Additional Concessions",        "$__________"),
        ("Inspection Item Agreed Repair", "__________________________________________"),
        ("Responsible Party for Repair",  "__________"),
        ("Completion Deadline",           "__________"),
    ]
    for label, placeholder in fields:
        doc.add_paragraph(f"{label}: {placeholder}")

    doc.add_paragraph("")
    doc.add_heading("SIGNATURES", 1)
    doc.add_paragraph("Buyer Signature: __________________________  Date: __________")
    doc.add_paragraph("Buyer Signature: __________________________  Date: __________")
    doc.add_paragraph("Seller Signature: _________________________  Date: __________")
    doc.add_paragraph("Seller Signature: _________________________  Date: __________")
    doc.add_paragraph("Agent: ___________________________________  License: __________")

    path = PROBLEM / "addendum_empty_fields.docx"
    doc.save(str(path))
    print(f"  Created: {path.relative_to(BASE)}")


def make_property_photo_duplicate(source_path: Path):
    """problematic/property_photo_duplicate.jpg — byte-for-byte copy of normal photo"""
    path = PROBLEM / "property_photo_duplicate.jpg"
    shutil.copy2(str(source_path), str(path))
    print(f"  Created: {path.relative_to(BASE)}  (SHA-256 == normal/property_photo.jpg)")


def make_hoa_missing_fields_csv():
    """problematic/hoa_missing_fields.csv — 6 of 12 data fields intentionally blank"""
    rows = [
        ["Field",                        "Value"],
        ["HOA Name",                     HOA_NAME],
        ["Property Address",             PROPERTY_ADDRESS],
        ["Parcel Number",                ""],           # blank
        ["Current Owner",                SELLER],
        ["Buyer Name",                   ""],           # blank
        ["Monthly Dues",                 ""],           # blank
        ["Capital Reserve Contribution", ""],           # blank
        ["Current Balance Due",          "$0.00"],
        ["Special Assessments Pending",  ""],           # blank
        ["Transfer Fee",                 ""],           # blank
        ["HOA Contact Phone",            "(704) 555-0182"],
        ["HOA Contact Email",            "info@pinewoodestates-hoa.com"],
    ]
    path = PROBLEM / "hoa_missing_fields.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"  Created: {path.relative_to(BASE)}")


def make_amendment_no_date():
    """problematic/amendment_no_date.txt — amendment with no date and no legal party names"""
    content = """PURCHASE AGREEMENT AMENDMENT

This amendment modifies the terms of the original purchase agreement for the above-referenced
property. The parties have agreed to the following modifications:

1. The purchase price shall be adjusted to reflect the agreed-upon repair credits.

2. The closing contingency period is hereby extended per mutual written agreement.

3. The seller agrees to address the following inspection items prior to closing:
   - HVAC service and certification
   - Repair of garage door spring mechanism
   - Correction of GFCI outlet in master bathroom
   - Replacement of missing attic insulation (R-30 minimum)

4. All other terms and conditions of the original Purchase Agreement remain in full force
   and effect and are not modified by this Amendment.

IN WITNESS WHEREOF, the parties have executed this Amendment as of the date last signed below.

Buyer: ___________________________    Date: _______________

Buyer: ___________________________    Date: _______________

Seller: __________________________    Date: _______________

Seller: __________________________    Date: _______________

NOTE: This document does not contain an effective date, buyer names, or seller names.
Parties must re-sign with complete information before this amendment is valid.
"""
    path = PROBLEM / "amendment_no_date.txt"
    path.write_text(content, encoding="utf-8")
    print(f"  Created: {path.relative_to(BASE)}")


def make_conflicting_data_json():
    """problematic/lender_data_conflicts.json — 3 conflicts in one JSON (price, name, date)"""
    data = {
        "transaction_id": "CC-2026-04-7421",
        "_note": "CONFLICT: purchase_price, buyer_name, and closing_date all differ from PA",
        "property": {
            "address": PROPERTY_ADDRESS,
            "parcel": PARCEL,
        },
        "buyers": [
            {
                "name": BUYER_WRONG,          # "Jim Mitchell" — conflicts with PA buyer name
                "email": "james.mitchell@email.com",
            },
        ],
        "financial": {
            "purchase_price": 345000.00,      # $345,000 — conflicts with PA $342,500
            "loan_amount": 314050.00,
            "cash_to_close": 34892.50,
            "seller_credits": 5000.00,        # $5,000 — conflicts with PA $3,500
            "closing_date": CLOSING_DATE_WRONG,  # July 1, 2026 — conflicts with PA
        },
        "lender": {
            "name": LENDER,
        },
        "conflicts_expected": [
            {"rule": "PA-001", "field": "buyer_name",     "expected": BUYER_CORRECT, "found": BUYER_WRONG},
            {"rule": "PA-003", "field": "purchase_price", "expected": 342500.00,     "found": 345000.00},
            {"rule": "PA-005", "field": "closing_date",   "expected": CLOSING_DATE_CORRECT, "found": CLOSING_DATE_WRONG},
            {"rule": "CD-003", "field": "seller_credits", "expected": 3500.00,       "found": 5000.00},
        ],
    }
    path = PROBLEM / "lender_data_conflicts.json"
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    print(f"  Created: {path.relative_to(BASE)}")


def make_scanned_page_image():
    """problematic/inspection_report_scan.png — Pillow-rendered text page (OCR target)"""
    W, H = 850, 1100
    img = Image.new("RGB", (W, H), color=(248, 246, 240))  # aged paper
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arial.ttf", 22)
        font_body  = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 13)
    except Exception:
        font_title = ImageFont.load_default()
        font_body  = font_title
        font_small = font_title

    y = 60
    def line(text, font=None, color=(30, 30, 30), indent=60):
        nonlocal y
        draw.text((indent, y), text, fill=color, font=font or font_body)
        y += 28

    line("HOME INSPECTION REPORT",                    font=font_title, color=(10, 10, 80))
    y += 10
    line(f"Property: {PROPERTY_ADDRESS}",             font=font_body)
    line(f"Inspection Date: May 3, 2026",             font=font_body)
    line(f"Inspector: Brian K. Walters, NC Lic #3841",font=font_body)
    line(f"Client: {BUYER_CORRECT}",                  font=font_body)
    y += 15
    line("─" * 70,                                    font=font_small)
    y += 5
    line("SUMMARY OF FINDINGS",                       font=font_title, color=(10, 10, 80))
    y += 8

    findings = [
        ("STRUCTURAL",       "Foundation appears sound. No evidence of settlement or cracking."),
        ("ROOF",             "Asphalt shingles approx. 8 years old. Minor granule loss on south slope."),
        ("ELECTRICAL",       "Panel updated 2018. GFCI missing in master bath — recommend correction."),
        ("PLUMBING",         "All fixtures functional. Water heater 2020, in good condition."),
        ("HVAC",             "HVAC system 2016. Last serviced unknown — recommend annual service."),
        ("INSULATION",       "Attic insulation R-19 observed in sections, R-30 recommended."),
        ("GARAGE",           "Garage door spring showing wear — recommend replacement before close."),
        ("EXTERIOR",         "Caulking around windows needs refreshing. No water intrusion noted."),
        ("INTERIOR",         "Minor cosmetic issues only. No structural concerns."),
    ]
    for category, desc in findings:
        line(f"• {category}: {desc}", font=font_small, indent=70)
        y += 4

    y += 15
    line("ITEMS REQUIRING ATTENTION",                 font=font_title, color=(140, 40, 40))
    y += 8
    attention = [
        "1. Install GFCI outlet in master bathroom (Code required)",
        "2. Service HVAC system — provide documentation",
        "3. Replace garage door spring (safety issue)",
        "4. Add attic insulation to minimum R-30",
    ]
    for a in attention:
        line(a, font=font_body, color=(140, 40, 40))

    y += 20
    line("─" * 70,                                    font=font_small)
    line("Inspector Signature: _________________________",  font=font_body)
    line(f"Date: May 3, 2026",                         font=font_body)

    # Slight rotation to simulate scan skew
    img = img.rotate(0.8, expand=False, fillcolor=(248, 246, 240))

    # Add slight noise / grain
    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))

    path = PROBLEM / "inspection_report_scan.png"
    img.save(str(path), "PNG")
    print(f"  Created: {path.relative_to(BASE)}")


# =============================================================================
# SECTION E — ZIP ARCHIVES
# =============================================================================

def make_clean_zip():
    """archives/closing_package_clean.zip — all normal/ files flat"""
    path = ARCHIVES / "closing_package_clean.zip"
    with zipfile.ZipFile(str(path), "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(NORMAL.iterdir()):
            if f.is_file():
                zf.write(str(f), f.name)
    print(f"  Created: {path.relative_to(BASE)}  ({len(list(NORMAL.iterdir()))} files)")


def make_mixed_zip():
    """archives/closing_package_mixed.zip — normal/ + selected problematic/ in subfolders"""
    path = ARCHIVES / "closing_package_mixed.zip"
    problem_files = [
        PROBLEM / "cd_amended_price_conflict.pdf",
        PROBLEM / "pa_signed_name_mismatch.docx",
        PROBLEM / "loan_estimate_date_conflict.pdf",
        PROBLEM / "closing_costs_bad_math.xlsx",
        PROBLEM / "lender_data_conflicts.json",
        PROBLEM / "amendment_no_date.txt",
    ]
    with zipfile.ZipFile(str(path), "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(NORMAL.iterdir()):
            if f.is_file():
                zf.write(str(f), f"normal/{f.name}")
        for f in problem_files:
            if f.is_file():
                zf.write(str(f), f"problematic/{f.name}")
    print(f"  Created: {path.relative_to(BASE)}")


def make_nested_zip():
    """archives/nested_archive.zip — outer ZIP containing inner_package.zip + __MACOSX entries"""
    # 1. Build the inner ZIP in memory
    inner_buf = io.BytesIO()
    with zipfile.ZipFile(inner_buf, "w", zipfile.ZIP_DEFLATED) as inner_zf:
        for f in [NORMAL / "purchase_agreement.pdf", NORMAL / "closing_disclosure.pdf"]:
            if f.is_file():
                inner_zf.write(str(f), f.name)
    inner_bytes = inner_buf.getvalue()

    # 2. Build the outer ZIP
    path = ARCHIVES / "nested_archive.zip"
    with zipfile.ZipFile(str(path), "w", zipfile.ZIP_DEFLATED) as zf:
        # Normal files at root level
        zf.write(str(NORMAL / "wire_instructions.txt"), "wire_instructions.txt")
        zf.write(str(NORMAL / "buyer_info.json"), "buyer_info.json")

        # Inner ZIP
        zf.writestr("inner_package.zip", inner_bytes)

        # __MACOSX metadata entries (should be skipped by zip_handler)
        zf.writestr("__MACOSX/._purchase_agreement.pdf", b"\x00\x05\x16\x07Apple")
        zf.writestr("__MACOSX/._buyer_info.json",        b"\x00\x05\x16\x07Apple")
        zf.writestr(".DS_Store",                         b"\x00\x00\x00\x01Bud1")

    print(f"  Created: {path.relative_to(BASE)}  (inner ZIP + __MACOSX entries)")


# =============================================================================
# SECTION F — MANIFEST
# =============================================================================

def make_manifest():
    """manifest.json — ground truth for integration test assertions"""
    manifest = {
        "transaction": {
            "property_address":  PROPERTY_ADDRESS,
            "parcel":            PARCEL,
            "buyer_correct":     BUYER_CORRECT,
            "buyer_wrong":       BUYER_WRONG,
            "seller":            SELLER,
            "purchase_price_correct": PURCHASE_PRICE_CORRECT,
            "purchase_price_wrong":   PURCHASE_PRICE_WRONG,
            "closing_date_correct":   CLOSING_DATE_CORRECT,
            "closing_date_wrong":     CLOSING_DATE_WRONG,
            "seller_credits_correct": SELLER_CREDITS_CORRECT,
            "seller_credits_wrong":   SELLER_CREDITS_WRONG,
            "loan_amount":       LOAN_AMOUNT,
            "cash_to_close":     CASH_TO_CLOSE,
            "lender":            LENDER,
        },
        "normal": {
            "purchase_agreement.pdf":   {"type": "purchase_agreement", "issues": []},
            "closing_disclosure.pdf":   {"type": "closing_disclosure", "issues": []},
            "title_commitment.pdf":     {"type": "title_commitment",   "issues": []},
            "loan_note.pdf":            {"type": "loan_note",          "issues": []},
            "insurance_binder.pdf":     {"type": "insurance_binder",   "issues": []},
            "hoa_certificate.pdf":      {"type": "hoa_document",       "issues": []},
            "wire_instructions.txt":    {"type": "wire_instructions",  "issues": []},
            "buyer_info.json":          {"type": "other",              "issues": []},
            "closing_costs.xlsx":       {"type": "other",              "issues": []},
            "settlement_summary.csv":   {"type": "other",              "issues": []},
            "property_photo.jpg":       {"type": "other",              "issues": [], "extraction": "ocr-vision"},
        },
        "problematic": {
            "cd_amended_price_conflict.pdf": {
                "type": "closing_disclosure",
                "issues": ["price conflict: $345,000 vs PA $342,500", "credits conflict: $5,000 vs PA $3,500"],
                "expected_rules_triggered": ["PA-003", "CD-003"],
            },
            "pa_signed_name_mismatch.docx": {
                "type": "purchase_agreement",
                "issues": ["buyer signed as 'Jim Mitchell' vs 'James R. Mitchell'"],
                "expected_rules_triggered": ["PA-001"],
            },
            "loan_estimate_date_conflict.pdf": {
                "type": "other",
                "issues": ["closing date July 1, 2026 vs PA June 15, 2026"],
                "expected_rules_triggered": ["PA-005"],
            },
            "corrupt_inspection_report.pdf": {
                "type": "other",
                "issues": ["corrupt PDF — valid header, invalid content streams"],
                "expected_pipeline_behavior": "extraction_method=failed, warning added",
            },
            "survey_scan_lowres.jpg": {
                "type": "survey",
                "issues": ["low resolution scan, heavy blur — OCR quality degraded"],
                "expected_pipeline_behavior": "extraction_method=ocr-vision",
            },
            "title_partial_missing_page.pdf": {
                "type": "title_commitment",
                "issues": ["document truncated — page 2 of 4 missing"],
                "expected_pipeline_behavior": "extraction partial, warning added",
            },
            "closing_costs_bad_math.xlsx": {
                "type": "other",
                "issues": ["TOTAL $5,436.50 does not equal sum of line items $4,186.50 (diff $1,250.00)"],
                "expected_rules_triggered": [],
                "note": "arithmetic error visible in spreadsheet data",
            },
            "addendum_empty_fields.docx": {
                "type": "purchase_agreement",
                "issues": ["11 unfilled placeholder fields (__________)"],
                "expected_pipeline_behavior": "extracted text contains placeholders",
            },
            "property_photo_duplicate.jpg": {
                "type": "other",
                "issues": ["exact byte-duplicate of normal/property_photo.jpg"],
                "expected_pipeline_behavior": "deduplication removes this file",
                "sha256_matches": "normal/property_photo.jpg",
            },
            "hoa_missing_fields.csv": {
                "type": "hoa_document",
                "issues": ["6 of 12 fields are blank"],
                "expected_pipeline_behavior": "extracted text has empty fields",
            },
            "amendment_no_date.txt": {
                "type": "other",
                "issues": ["no effective date", "no legal party names"],
                "expected_pipeline_behavior": "extraction succeeds, classification likely 'other'",
            },
            "lender_data_conflicts.json": {
                "type": "other",
                "issues": [
                    "purchase_price=$345,000 conflicts with PA $342,500",
                    "buyer='Jim Mitchell' conflicts with PA buyer name",
                    "closing_date='July 1, 2026' conflicts with PA",
                    "seller_credits=$5,000 conflicts with PA $3,500",
                ],
                "expected_rules_triggered": ["PA-001", "PA-003", "PA-005", "CD-003"],
            },
            "inspection_report_scan.png": {
                "type": "other",
                "issues": [],
                "expected_pipeline_behavior": "extraction_method=ocr-vision",
            },
        },
        "archives": {
            "closing_package_clean.zip": {
                "description": "All 11 normal/ documents in a flat ZIP",
                "expected_rules_triggered": [],
                "expected_file_count_after_expansion": 11,
            },
            "closing_package_mixed.zip": {
                "description": "normal/ folder + 6 problematic files in subfolders",
                "expected_rules_triggered": ["PA-001", "PA-003", "PA-005", "CD-003"],
                "expected_file_count_after_expansion": 17,
            },
            "nested_archive.zip": {
                "description": "Outer ZIP with inner_package.zip + __MACOSX entries to skip",
                "expected_pipeline_behavior": "recursive extraction; __MACOSX and .DS_Store skipped",
                "expected_file_count_after_expansion": 4,
            },
        },
    }
    path = BASE / "manifest.json"
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"  Created: {path.relative_to(BASE)}")


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("CloseCheck test dataset generator")
    print(f"Output: {DOCS}\n")
    _mkdir()

    print("[1/4] Normal documents (11 files):")
    make_purchase_agreement()
    make_closing_disclosure()
    make_title_commitment()
    make_loan_note()
    make_insurance_binder()
    make_hoa_certificate()
    make_wire_instructions()
    make_buyer_info_json()
    make_closing_costs_xlsx()
    make_settlement_summary_csv()
    photo_path = make_property_photo()

    print("\n[2/4] Problematic documents (13 files):")
    make_cd_price_conflict()
    make_pa_signed_name_mismatch()
    make_loan_estimate_date_conflict()
    make_corrupt_inspection_report()
    make_survey_scan_lowres()
    make_title_partial()
    make_closing_costs_bad_math()
    make_addendum_empty_fields()
    make_property_photo_duplicate(NORMAL / "property_photo.jpg")
    make_hoa_missing_fields_csv()
    make_amendment_no_date()
    make_conflicting_data_json()
    make_scanned_page_image()

    print("\n[3/4] ZIP archives (3 files):")
    make_clean_zip()
    make_mixed_zip()
    make_nested_zip()

    print("\n[4/4] Manifest:")
    make_manifest()

    # Summary
    n_normal   = len(list(NORMAL.iterdir()))
    n_problem  = len(list(PROBLEM.iterdir()))
    n_archives = len(list(ARCHIVES.iterdir()))
    print(f"\n{'─'*50}")
    print(f"  Normal documents:    {n_normal:3d}")
    print(f"  Problematic docs:    {n_problem:3d}")
    print(f"  ZIP archives:        {n_archives:3d}")
    print(f"  Total:               {n_normal + n_problem + n_archives:3d}")
    print(f"{'─'*50}")
    print(f"\nManifest: {BASE / 'manifest.json'}")
    print("\nDone.")


if __name__ == "__main__":
    main()
